import os
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from config import UPLOAD_DIR

class DocumentGenerator:
    """
    Generates professional legal documents using python-docx.
    """
    def __init__(self):
        pass

    def _get_output_path(self, user_id: int, doc_type: str) -> Path:
        """
        Creates output directory if it doesn't exist and returns file path.
        """
        out_dir = UPLOAD_DIR / str(user_id) / "generated"
        out_dir.mkdir(parents=True, exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        return out_dir / f"{doc_type}_{timestamp}.docx"

    def _add_header(self, doc, case_info: dict, title: str):
        """
        Adds title, client name, date, and case type to the document header.
        """
        # Title
        heading = doc.add_heading(title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"Date: {datetime.now().strftime('%d %B %Y')}")
        
        # Case Info Table
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Light Shading Accent 1'
        
        client_name = case_info.get("client_name", "N/A")
        case_type = case_info.get("case_type", "N/A")
        opposing = case_info.get("opposing_party_name", "N/A")
        court = case_info.get("court_name", "N/A")
        
        cells = table.rows[0].cells
        cells[0].text = "Client Name:"
        cells[1].text = client_name
        
        cells = table.rows[1].cells
        cells[0].text = "Case Type:"
        cells[1].text = str(case_type)
        
        cells = table.rows[2].cells
        cells[0].text = "Opposing Party:"
        cells[1].text = opposing
        
        cells = table.rows[3].cells
        cells[0].text = "Court:"
        cells[1].text = court
        
        doc.add_paragraph() # spacing

    def _add_section(self, doc, title: str, content: str, level: int = 2):
        """
        Adds a section with a heading to the document.
        """
        doc.add_heading(title, level=level)
        if content:
            doc.add_paragraph(content)

    def _add_references(self, doc, references: list[dict]):
        """
        Adds numbered references with Kanoon URLs to the document.
        """
        if not references:
            return
            
        doc.add_heading("References & Precedents", level=2)
        
        for i, ref in enumerate(references, 1):
            title = ref.get("title", "Unknown Title")
            p = doc.add_paragraph(style='List Number')
            runner = p.add_run(f"{title}")
            runner.bold = True
            
            court = ref.get("court")
            date = ref.get("date")
            if court or date:
                p.add_run(f" ({court or ''} {date or ''})".strip())
                
            kanoon_id = ref.get("kanoon_doc_id")
            if kanoon_id:
                p.add_run(f"\nSource: https://indiankanoon.org/doc/{kanoon_id}/")
            
            snippet = ref.get("snippet")
            if snippet:
                snippet_para = doc.add_paragraph(snippet)
                snippet_para.paragraph_format.left_indent = Inches(0.5)
                # Italicize snippet
                for run in snippet_para.runs:
                    run.italic = True

    def _add_disclaimer(self, doc):
        """
        Adds footer disclaimer about AI assistance.
        """
        doc.add_paragraph()
        p = doc.add_paragraph("Disclaimer: This document was generated with AI assistance. It is intended for review and finalization by a qualified legal professional before submission or official use.")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(8)
            run.font.italic = True

    def generate_case_brief(self, case_info: dict, sections: dict, references: list[dict]) -> str:
        """
        Generates a Case Brief DOCX file.
        """
        doc = Document()
        self._add_header(doc, case_info, "CASE BRIEF")
        
        # Standard sections for a case brief
        for key, title in [
            ("facts", "1. Material Facts"),
            ("issues", "2. Issues Presented"),
            ("applicable_law", "3. Applicable Law"),
            ("precedent_analysis", "4. Analysis of Precedents"),
            ("arguments", "5. Arguments"),
            ("prayer", "6. Conclusion / Prayer")
        ]:
            if key in sections:
                self._add_section(doc, title, sections[key])
        
        self._add_references(doc, references)
        self._add_disclaimer(doc)
        
        user_id = case_info.get("user_id", 0)
        out_path = self._get_output_path(user_id, "case_brief")
        doc.save(out_path)
        return str(out_path)

    def generate_legal_notice(self, case_info: dict, sections: dict, references: list[dict]) -> str:
        """
        Generates a Legal Notice DOCX file.
        """
        doc = Document()
        self._add_header(doc, case_info, "LEGAL NOTICE")
        
        for key, title in [
            ("introduction", "1. Introduction"),
            ("facts", "2. Facts and Background"),
            ("breach", "3. Grievance / Breach"),
            ("demands", "4. Demands"),
            ("consequences", "5. Consequences of Non-Compliance")
        ]:
            if key in sections:
                self._add_section(doc, title, sections[key])
                
        self._add_references(doc, references)
        self._add_disclaimer(doc)
        
        user_id = case_info.get("user_id", 0)
        out_path = self._get_output_path(user_id, "legal_notice")
        doc.save(out_path)
        return str(out_path)

    def generate_case_analysis(self, case_info: dict, sections: dict, references: list[dict]) -> str:
        """
        Generates a Case Analysis DOCX file.
        """
        doc = Document()
        self._add_header(doc, case_info, "CASE ANALYSIS")
        
        for key, title in [
            ("summary", "1. Executive Summary"),
            ("strengths", "2. Strengths of the Case"),
            ("weaknesses", "3. Weaknesses and Risks"),
            ("strategy", "4. Proposed Strategy")
        ]:
            if key in sections:
                self._add_section(doc, title, sections[key])
                
        self._add_references(doc, references)
        self._add_disclaimer(doc)
        
        user_id = case_info.get("user_id", 0)
        out_path = self._get_output_path(user_id, "case_analysis")
        doc.save(out_path)
        return str(out_path)
